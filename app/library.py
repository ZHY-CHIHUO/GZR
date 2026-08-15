# -*- coding: utf-8 -*-
"""阅读库：原版小说目录/章节全文、插图版PDF、人祖传、资料合集HTML。"""
import html
import os
import re
from pathlib import Path

BASE = Path(__file__).resolve().parent.parent
NOVEL_ROOT = BASE.parent / "gu-zhen-ren"
PDF_ROOT = BASE.parent / "gu_zhen_ren_pdf"
LORE_DOCX = BASE.parent / "gu-zhenren-lore" / "蛊真人资料合集.docx"


_CN = {"零": 0, "一": 1, "二": 2, "三": 3, "四": 4, "五": 5, "六": 6, "七": 7, "八": 8, "九": 9}


def _cn_vol_num(s: str):
    """提取 第N卷/章/节 的汉字数字并转整数（支持 一~九、十、几十、几十几）。"""
    m = re.search(r"第([零一二三四五六七八九十两]+)(?:卷|章|节)", s)
    if not m:
        return 0
    txt = m.group(1)
    if "十" not in txt:
        return _CN.get(txt, 0)
    a, _, b = txt.partition("十")
    return (_CN.get(a, 1) if a else 1) * 10 + _CN.get(b, 0)


def _natural_key(s: str):
    # 汉字数字卷/章序优先（“第二卷” < “第三卷”），其余按数字/字符兜底
    return (_cn_vol_num(s), [int(t) if t.isdigit() else t for t in re.split(r"(\d+)", s)])


def _first_title(lines):
    for ln in lines:
        if ln.strip() and not re.fullmatch(r"[=\-—\s·]+", ln.strip()):
            return ln.strip()
    return lines[0].strip() if lines else ""


def novel_volumes():
    """返回 [{name, chapters: [{n, title, file}]}]"""
    vols = []
    if not NOVEL_ROOT.is_dir():
        return vols
    for vdir in sorted(os.listdir(NOVEL_ROOT), key=_natural_key):
        vpath = NOVEL_ROOT / vdir
        if not vpath.is_dir():
            continue
        chapters = []
        for fn in sorted(os.listdir(vpath), key=_natural_key):
            if not fn.endswith(".txt"):
                continue
            m = re.search(r"第(\d+)章", fn)
            n = int(m.group(1)) if m else 0
            try:
                with open(vpath / fn, encoding="utf-8", errors="replace") as f:
                    lines = [l.strip() for l in f.read().splitlines() if l.strip()]
            except OSError:
                lines = []
            chapters.append({"n": n, "title": _first_title(lines), "file": fn})
        if chapters:
            vols.append({"name": vdir, "chapters": chapters})
    return vols


def chapter_text(vol: str, chapter: int):
    """返回 {vol, chapter, title, text, path} 或 None"""
    vdir = NOVEL_ROOT / vol
    if not vdir.is_dir():
        return None
    if chapter == 0:
        cand = vdir / "序.txt"
        fname = "序.txt"
    else:
        fname = f"第{chapter}章.txt"
        cand = vdir / fname
    if not cand.is_file():
        # 容错：按文件名匹配 第N章
        for fn in os.listdir(vdir):
            m = re.search(rf"第{chapter}章", fn)
            if m:
                cand = vdir / fn
                fname = fn
                break
        else:
            return None
    try:
        with open(cand, encoding="utf-8", errors="replace") as f:
            raw = f.read()
    except OSError:
        return None
    lines = [l.strip() for l in raw.splitlines() if l.strip()]
    if not lines:
        return None
    return {
        "vol": vol, "chapter": chapter, "file": fname,
        "title": _first_title(lines),
        "text": "\n".join(lines[1:]),
        "path": str(cand),
    }


_CN_DIGIT = {"零": 0, "一": 1, "二": 2, "三": 3, "四": 4, "五": 5, "六": 6, "七": 7, "八": 8, "九": 9}


def _cn_to_int(s):
    """汉字数字 -> 整数（支持 十/百/千，如 一百零一 -> 101）。"""
    s = s.replace("零", "")
    total, cur = 0, 0
    for c in s:
        if c in _CN_DIGIT:
            cur = _CN_DIGIT[c]
        elif c == "十":
            total += (cur or 1) * 10
            cur = 0
        elif c == "百":
            total += (cur or 1) * 100
            cur = 0
        elif c == "千":
            total += (cur or 1) * 1000
            cur = 0
    return total + cur


_pdf_toc_cache = None


def _combined_flat_toc():
    """合订本(1.1)目录：用原版卷结构重建，章节通过标题匹配到合订本页码。
    返回扁平 [{title, page, depth}]（卷=0，章=1）。"""
    from pypdf import PdfReader
    path = PDF_ROOT / "蛊真人" / "蛊无删减插图版（1.1版）.pdf"
    if not path.is_file():
        return []
    r = PdfReader(str(path))
    page_map = {}
    front = []   # 前言：制作说明/供稿/阅读方法/序等
    extras = []  # 番外篇等
    vols = novel_volumes()
    vol_words = set()
    for v in vols:
        vol_words.add(v["name"])
        if "：" in v["name"]:
            vol_words.add(v["name"].split("：")[-1])
    for it in r.outline:
        if isinstance(it, list):
            continue
        t = str(it.title or "").strip()
        if not t:
            continue
        try:
            pg = int(r.get_destination_page_number(it)) + 1
        except Exception:
            pg = None
        m = re.match(r"^第[零一二三四五六七八九十百千两]+节[：:]\s*(.*)$", t)
        if m:
            page_map[m.group(1).strip()] = pg
        elif t in vol_words or re.match(r"^第[零一二三四五六七八九十两]+卷", t):
            continue  # 跳过纯卷名/残缺卷级书签
        elif t.startswith("番外"):
            t2 = re.sub(r"^(番外篇第[零一二三四五六七八九十百千两]+章)(.*)$", r"\1 · \2", t)
            extras.append((t2, pg))
        else:
            front.append((t, pg))
    out = []
    if front:
        out.append({"title": "前言", "page": front[0][1], "depth": 0})
        for t, pg in front:
            out.append({"title": t, "page": pg, "depth": 1})
    for v in vols:
        items = []
        for c in v["chapters"]:
            t = re.sub(r"^第[零一二三四五六七八九十百千两]+节[：:]\s*", "", c["title"]).strip()
            pg = page_map.get(t)
            if pg:
                items.append((c["n"], t, pg))
        if items:
            out.append({"title": v["name"], "page": items[0][2], "depth": 0})
            for n, t, pg in items:
                out.append({"title": f"第{n}章 · {t}", "page": pg, "depth": 1})
    if extras:
        out.append({"title": "番外", "page": extras[0][1], "depth": 0})
        for t, pg in extras:
            out.append({"title": t, "page": pg, "depth": 1})
    return out


def _normalize_toc(fn, items):
    """按标题模式重建/过滤书签层级：
    - 人祖传：只保留「人祖传（N）——」章标题
    - 合订本(1.1)：序/第X卷 = 一级(depth0)，第X节 = 二级(depth1)
    """
    if "人祖传" in fn:
        return [it for it in items if re.match(r"^人祖传（[零一二三四五六七八九十两]+）——", it["title"])]
    if "1.1" in fn:
        return _combined_flat_toc()
    # 分卷 PDF：只保留 第X节，并统一为 第N章 · 标题
    out = []
    for it in items:
        m = re.match(r"^第([零一二三四五六七八九十百千两]+)节[：:]\s*(.*)$", it["title"])
        if m:
            n = _cn_to_int(m.group(1))
            out.append(dict(it, title=f"第{n}章 · {m.group(2).strip()}", depth=0))
    return out


def pdf_toc():
    """提取所有 PDF 的书签目录，带磁盘缓存。返回 {文件名: [{title, page, depth}]}"""
    global _pdf_toc_cache
    if _pdf_toc_cache is not None:
        return _pdf_toc_cache
    import json as _json
    from app.config import BASE as _BASE
    cache = _BASE / "data" / "pdf_toc_cache.json"
    if cache.is_file():
        try:
            _pdf_toc_cache = _json.loads(cache.read_text(encoding="utf-8"))
            return _pdf_toc_cache
        except Exception:
            pass
    from pypdf import PdfReader
    out = {}
    for gdir in sorted(os.listdir(PDF_ROOT)):
        gp = PDF_ROOT / gdir
        if not gp.is_dir():
            continue
        for fn in sorted(os.listdir(gp)):
            if not fn.lower().endswith(".pdf"):
                continue
            path = gp / fn
            try:
                r = PdfReader(str(path))
                items = []
                def walk(lst, depth):
                    for it in lst:
                        if isinstance(it, list):
                            walk(it, depth + 1)
                            continue
                        title = str(it.title or "").strip()
                        page = None
                        try:
                            page = int(r.get_destination_page_number(it)) + 1
                        except Exception:
                            page = None
                        if title:
                            items.append({"title": title, "page": page, "depth": depth})
                walk(r.outline or [], 0)
                if items:
                    out[fn] = _normalize_toc(fn, items)
            except Exception as e:
                print(f"[pdf_toc] 解析失败 {fn}: {str(e)[:60]}")
    _pdf_toc_cache = out
    try:
        cache.write_text(_json.dumps(out, ensure_ascii=False), encoding="utf-8")
    except Exception:
        pass
    return out


def pdf_files():
    """返回 [{group, name, path, url}]"""
    out = []
    if not PDF_ROOT.is_dir():
        return out
    groups = {"蛊真人": "插图版", "人祖传": "人祖传"}
    for gdir in sorted(os.listdir(PDF_ROOT), key=_natural_key):
        gp = PDF_ROOT / gdir
        if not gp.is_dir():
            continue
        for fn in sorted(os.listdir(gp), key=_natural_key):
            if fn.lower().endswith(".pdf"):
                rel = f"{gdir}/{fn}"
                out.append({
                    "group": groups.get(gdir, gdir),
                    "name": fn,
                    "url": f"/files/{rel}",
                    "path": str(gp / fn),
                })
    return out


_lore_html_cache = None
_lore_structured_cache = None


def lore_structured():
    """资料合集结构化数据（供主应用统一 UI 渲染）：{title, toc:[{text,level,anchor}], paras:[{kind,level,text,anchor}]}
    磁盘缓存：首次生成后写入 data/lore_structured_cache.json，之后秒开（docx 解析约 12 秒）。"""
    global _lore_structured_cache
    if _lore_structured_cache is not None:
        return _lore_structured_cache
    from docx import Document
    import json as _json
    from app.config import DATA_DIR as _DD, BASE as _BASE
    _cache_file = _BASE / "data" / "lore_structured_cache.json"
    _toc_path = _DD / "lore_toc.json"
    _toc_mtime = _toc_path.stat().st_mtime if _toc_path.is_file() else 0
    if _cache_file.is_file() and _cache_file.stat().st_mtime >= _toc_mtime:
        try:
            _lore_structured_cache = _json.loads(_cache_file.read_text(encoding="utf-8"))
            return _lore_structured_cache
        except Exception:
            pass

    def _looks_body(t):
        if re.match(r"^\d+[.、]\s*", t):
            return True
        if len(t) > 30 or t.endswith(("。", "，", "！", "？", "；", "、", "：", ":")):
            return True
        return False

    def _heading_level(t):
        if _looks_body(t) or len(t) < 2 or len(t) > 24:
            return None
        if t.startswith(("【", "[", "（", "(")) or re.fullmatch(r"[\]\]）\)\-—=·\s]+", t):
            return None
        if "：" in t or ":" in t:
            return None
        return 2

    doc = Document(str(LORE_DOCX))
    raw = [(p.text.strip(), (p.style.name if p.style else "") or "")
           for p in doc.paragraphs if p.text and p.text.strip()]
    paras = [(t, s) for t, s in raw if not re.fullmatch(r"\[\d+\]\s*", t)]
    merged = []
    for t, s in paras:
        if (t.startswith("【") or t.startswith("】")) and merged:
            merged[-1] = (merged[-1][0] + t, merged[-1][1])
        else:
            merged.append((t, s))
    toc_items = None
    tp = _DD / "lore_toc.json"
    if tp.is_file():
        try:
            toc_items = _json.loads(tp.read_text(encoding="utf-8"))["items"]
        except Exception:
            toc_items = None

    toc, out = [], []
    idx = 0
    seen = set()
    for pi, (t, s) in enumerate(merged):
        lv = None
        txt = t
        if toc_items is not None and pi < len(toc_items):
            lv = int(toc_items[pi].get("level", 0))
            txt = str(toc_items[pi].get("text") or t)
        else:
            lv = _heading_level(t)
        if lv:
            idx += 1
            anchor = f"sec{idx}"
            if txt not in seen:
                seen.add(txt)
                toc.append({"text": txt, "level": min(lv, 3), "anchor": anchor})
            out.append({"kind": "h2", "level": min(lv, 3), "text": txt, "anchor": anchor})
        else:
            out.append({"kind": "p", "text": txt})
    _lore_structured_cache = {"title": "《蛊真人》资料合集", "toc": toc, "paras": out}
    try:
        _cache_file.write_text(_json.dumps(_lore_structured_cache, ensure_ascii=False), encoding="utf-8")
    except Exception:
        pass
    return _lore_structured_cache


def lore_html():
    """资料合集 docx -> HTML（带目录锚点），进程内缓存。"""
    global _lore_html_cache
    if _lore_html_cache is not None:
        return _lore_html_cache
    from docx import Document
    doc = Document(str(LORE_DOCX))
    raw = [(p.text.strip(), (p.style.name if p.style else "") or "")
           for p in doc.paragraphs if p.text and p.text.strip()]
    paras = [(t, s) for t, s in raw if not re.fullmatch(r"\[\d+\]\s*", t)]

    # 合并残段：以【/】开头的行是上一条目的尾注残段，拼回前一行
    merged = []
    for t, s in paras:
        if (t.startswith("【") or t.startswith("】")) and merged:
            merged[-1] = (merged[-1][0] + t, merged[-1][1])
        else:
            merged.append((t, s))

    # 一级大节名（docx 无样式层级，用内容启发式分级）
    L1_KEYWORDS = ("蛊虫百科", "百科内容", "作品简介", "作品目录", "作品设定", "背景设定",
                   "世界观", "修行体系", "流派", "人物图鉴", "势力分布", "仙蛊屋全集", "仙蛊屋",
                   "灾劫资料", "杀招体系", "荒兽", "人祖传", "尊者", "语录", "金句", "访谈",
                   "资料统计", "蛊仙数据", "奇蛊榜", "仙蛊榜", "魔蛊榜", "境界", "蛊师相关")

    def looks_like_body(t):
        """内容像正文（编号行/长句/句末标点）的行，即使带标题样式也不算标题。"""
        if re.match(r"^\d+[\.、]\s*", t):
            return True
        if len(t) > 30:
            return True
        if t.endswith(("。", "，", "！", "？", "；", "、", "：", ":")):
            return True
        return False

    def heading_level(t, style):
        """返回 1/2/None（None=非标题）。样式优先，内容校验兜底。"""
        if style:
            if "Heading 1" in style or "标题 1" in style:
                return None if looks_like_body(t) else 1
            if "Heading" in style or "标题" in style:
                return None if looks_like_body(t) else 2
        if looks_like_body(t):
            return None
        if len(t) < 2 or len(t) > 24:
            return None
        if t.startswith(("【", "[", "（", "(")):
            return None
        if re.fullmatch(r"[\]\]）\)\-—=·\s]+", t):
            return None
        if "：" in t or ":" in t:
            return None
        if any(k in t for k in L1_KEYWORDS) and len(t) <= 14:
            return 1
        return 2

    # 优先使用 AI 审核的目录（scripts/ai_toc.py 生成）
    import json as _json
    from app.config import DATA_DIR as _DD
    toc_items = None
    _toc_path = _DD / "lore_toc.json"
    if _toc_path.is_file():
        try:
            toc_items = _json.loads(_toc_path.read_text(encoding="utf-8"))["items"]
        except Exception:
            toc_items = None

    parts = ["<h1>《蛊真人》资料合集</h1>"]
    toc = []
    idx = 0
    seen_toc = set()
    toc_stack = []  # (level, 是否已生成该级 summary)
    for pi, (p, style) in enumerate(merged):
        lv = None
        txt = p
        if toc_items is not None and pi < len(toc_items):
            lv = int(toc_items[pi].get("level", 0))
            txt = str(toc_items[pi].get("text") or p)
        else:
            lv = heading_level(p, style)
        e = html.escape(txt)
        if lv:
            idx += 1
            anchor = f"sec{idx}"
            cls = f"l{min(lv, 3)}"
            parts.append(f'<h2 class="{cls}" id="{anchor}">{e}</h2>')
            if e in seen_toc:
                continue
            seen_toc.add(e)
            # 关闭比当前级别深的 open details
            while toc_stack and toc_stack[-1] >= lv:
                toc_stack.pop()
                toc.append("</details>")
            if lv == 3:
                toc.append(f'<a class="toc-l3" href="#{anchor}">{e}</a>')
            else:
                toc.append(f'<details open class="toc-grp"><summary><a href="#{anchor}">{e}</a></summary>')
                toc_stack.append(lv)
        else:
            parts.append(f"<p>{e}</p>")
    while toc_stack:
        toc_stack.pop()
        toc.append("</details>")
    _lore_html_cache = (
        "<style>body{font-family:'PingFang SC','Microsoft YaHei',sans-serif;line-height:1.9;margin:0;background:#f5f3ee}"
        "#toc{position:fixed;left:0;top:0;bottom:0;width:240px;overflow-y:auto;background:#fffdf7;border-right:1px solid #e5e0d6;padding:14px 10px;font-size:12px;z-index:5}"
        "#toc h3{font-size:13px;margin:0 0 8px 6px;color:#7a5c3e;cursor:pointer;user-select:none}"
        "#toc a{display:block;color:#7a5c3e;text-decoration:none;margin:2px 0;line-height:1.6;overflow:hidden;text-overflow:ellipsis;white-space:nowrap}"
        "#toc a:hover{background:#f0e9da;border-radius:4px}"
        "#toc summary{list-style:none;cursor:pointer;font-weight:700;color:#5a4530;margin-top:6px;user-select:none}"
        "#toc summary::-webkit-details-marker{display:none}"
        "#toc summary::before{content:'▾ ';color:#b08d57}"
        "#toc details:not([open])>summary::before{content:'▸ '}"
        "#toc summary a{display:inline;margin:0;font-weight:700}"
        "#toc .toc-grp2 summary{font-weight:600;color:#7a5c3e;margin-top:2px;padding-left:10px}"
        "#toc .toc-grp2 summary::before{content:'▾ '}"
        "#toc .toc-grp2 a{font-weight:400;display:inline}"
        "#toc .toc-grp{padding-left:2px}"
        "#toc .toc-grp2{padding-left:10px}"
        "#toc .toc-l3{padding-left:24px;color:#9a9284;font-size:11px}"
        "#content{margin-left:260px;max-width:860px;padding:24px 32px 60px;background:#fff;min-height:100vh}"
        "#toc .toc-top{display:flex;justify-content:space-between;align-items:center;font-size:13px;color:#7a5c3e;margin:0 2px 8px}"
        "#toc .toc-top button{border:1px solid #e5e0d6;background:#fff;border-radius:6px;padding:1px 10px;cursor:pointer;font-size:12px}"
        "#toc-show{position:sticky;top:0;z-index:9;border:1px solid #e5e0d6;background:#fffdf7;border-radius:8px;padding:5px 14px;font-size:13px;cursor:pointer;color:#7a5c3e;margin-bottom:10px}"
        "body.toc-hidden #toc{width:30px;padding:10px 3px;background:rgba(255,253,247,.5);transition:background .15s}"
        "body.toc-hidden #toc:hover{background:rgba(255,253,247,1)}"
        "body.toc-hidden #tocbody{display:none}"
        "body.toc-hidden #toc .toc-title{display:none}"
        "body.toc-hidden #toc .toc-top{justify-content:center}"
        "body.toc-hidden #toc .toc-top button{transform:rotate(180deg)}"
        "body.toc-hidden #content{margin-left:36px}"
        "h1{font-size:22px}h2.l1{font-size:18px;margin-top:32px;border-left:4px solid #7a5c3e;padding-left:10px}"
        "h2.l2{font-size:15px;margin-top:22px;color:#5a4530;border-left:3px solid #b08d57;padding-left:8px}"
        "h2.l3{font-size:14px;margin-top:16px;color:#7a5c3e;padding-left:6px}"
        "p{margin:8px 0;color:#333}</style>"
        '<div id="toc"><div class="toc-top"><span class="toc-title">📑 目录</span><button onclick="toggleTocSide()" title="收起/展开目录">«</button></div><div id="tocbody">' + "".join(toc) + "</div></div>"
        + '<div id="content">' + "".join(parts) + "</div>"
        + "<script>function toggleTocSide(){document.body.classList.toggle('toc-hidden');}</script>"
    )
    return _lore_html_cache
